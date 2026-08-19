"""Build a question paper .docx in the format-2 table style.

A pipeline builder: takes the universal questions JSON (the same schema
run_pipeline.py validates) and emits the format-2 layout. Unicode Hindi in the
JSON is converted to Kruti Dev 010 runs (English words stay Latin) via
unicode_to_krutidev_runs — the same text path build_paper.py uses.

Each question becomes one 8-row × 3-col table:

  Row 0  Question   | <stem, colspan 2>
  Row 1  Type       | multiple_choice (colspan 2)
  Row 2  Option     | <opt A>  | incorrect
  Row 3  Option     | <opt B>  | incorrect
  Row 4  Option     | <opt C>  | incorrect
  Row 5  Option     | <opt D>  | incorrect
  Row 6  Solution   | (blank, colspan 2)
  Row 7  Marks      | 1        | 0.25

All options marked "incorrect" — client deletes "in" from the correct one.

Usage:
  python build_paper_format2.py <questions.json> [out.docx]
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.oxml   import OxmlElement
from docx.oxml.ns import qn
from docx.shared  import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
from krutidev import DEFAULT_FONT, render_runs  # noqa: E402

BASE    = Path(__file__).resolve().parents[1]
# review/ is client-specific — see build_answer_key.py's CLIENT_DIR comment.
CLIENT_DIR = BASE.parent / "clients" / "target-academy"
OUT_DIR = CLIENT_DIR / "review" / "output"

LATIN = "Times New Roman"
_FONT = DEFAULT_FONT  # set by build(); _runs() reads it via render_runs()

# ── Exact measurements from format-2.docx XML ────────────────────────────────
TBL_W      = 9150
TBL_INDENT = 216
COL0_W     = 1680   # label column
COL1_W     = 3288   # content column
COL2_W     = 4182   # meta column
TBL_FILL   = "CED7E7"
BORDER_SZ  = 6
CELL_MAR   = 81
BODY_PT    = 12


# ── XML helpers ───────────────────────────────────────────────────────────────

def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), str(v))
    return e


def _border(side):
    return _el(f"w:{side}", **{
        "w:val": "single", "w:sz": str(BORDER_SZ),
        "w:space": "0", "w:color": "000000"
    })


def _apply_tcPr(tc, width, span=1):
    tcPr = OxmlElement("w:tcPr")
    tcPr.append(_el("w:tcW", **{"w:w": str(width), "w:type": "dxa"}))
    if span > 1:
        tcPr.append(_el("w:gridSpan", **{"w:val": str(span)}))
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        borders.append(_border(side))
    tcPr.append(borders)
    mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "right"):
        mar.append(_el(f"w:{side}", **{"w:w": str(CELL_MAR), "w:type": "dxa"}))
    tcPr.append(mar)
    tc.insert(0, tcPr)


def _apply_tblPr(tbl):
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblPr.append(_el("w:tblStyle",  **{"w:val": "a4"}))
    tblPr.append(_el("w:tblW",      **{"w:w": str(TBL_W), "w:type": "dxa"}))
    tblPr.append(_el("w:tblInd",    **{"w:w": str(TBL_INDENT), "w:type": "dxa"}))
    tblPr.append(_el("w:shd",       **{
        "w:val": "clear", "w:color": "auto", "w:fill": TBL_FILL
    }))
    tblPr.append(_el("w:tblLayout", **{"w:type": "fixed"}))
    tblPr.append(_el("w:tblLook",   **{
        "w:val": "0600", "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1"
    }))


def _apply_tblGrid(tbl):
    grid = OxmlElement("w:tblGrid")
    for w in (COL0_W, COL1_W, COL2_W):
        grid.append(_el("w:gridCol", **{"w:w": str(w)}))
    tbl._tbl.insert(1, grid)


def _fmt_para(para):
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.space_before = Pt(0)


def _add_run(para, text, bold=True, font=LATIN):
    r = para.add_run(text)
    r.bold           = bold
    r.font.name      = font
    r.font.size      = Pt(BODY_PT)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return r


def _fill_label(cell, label_text, width=COL0_W, span=1):
    """Fill a label cell (Question / Type / Option / Solution / Marks)."""
    _apply_tcPr(cell._tc, width, span=span)
    para = cell.paragraphs[0]
    _fmt_para(para)
    _add_run(para, label_text, bold=True, font=LATIN)


def _fill_runs(cell, runs, width, span=1, bold=False):
    """Fill a content cell from a list of (text, font) run tuples."""
    _apply_tcPr(cell._tc, width, span=span)
    # Remove the second cell when colspan=2
    if span == 2:
        tc2 = cell._tc.getnext()
        if tc2 is not None:
            tc2.getparent().remove(tc2)
    para = cell.paragraphs[0]
    _fmt_para(para)
    for text, font in runs:
        _add_run(para, text, bold=bold, font=font)


def _fill_plain(cell, text, width, span=1, bold=False, font=LATIN):
    """Fill a content cell with a plain Latin string."""
    _fill_runs(cell, [(text, font)], width, span=span, bold=bold)


def _runs(text):
    """Unicode string -> [(text, font)] runs: Hindi in the chosen Devanagari font,
    English in Latin. Font follows _FONT (set by build())."""
    runs, deva_font = render_runs(text, _FONT)
    return [(t, LATIN if is_latin else deva_font) for t, is_latin in runs]


def _no_split_rows(tbl):
    """Keep one question's whole table on a SINGLE page.

    Two things are needed and w:cantSplit alone is NOT enough:
      1. w:cantSplit on each row — stops a single row from breaking mid-row.
      2. w:keepNext on every paragraph of every row EXCEPT the last — this is
         what actually stops the TABLE from breaking BETWEEN rows across a page
         (the bug where Question/Type land on page 1 and the Options on page 2).
         keepNext chains each row to the next, so Word floats the entire table
         to the next page when it doesn't fit, rather than splitting it.
    """
    rows = tbl.rows
    last = len(rows) - 1
    for i, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(_el("w:cantSplit", **{"w:val": "true"}))
        if i == last:
            continue  # the last row has nothing after it to keep with
        for cell in row.cells:
            for para in cell.paragraphs:
                pPr = para._p.get_or_add_pPr()
                # avoid stacking duplicate keepNext if a cell is revisited
                if pPr.find(qn("w:keepNext")) is None:
                    pPr.append(_el("w:keepNext", **{"w:val": "true"}))


def add_question_table(doc, q, student_safe: bool = True):
    """Render one question as the format-2 table.

    student_safe (default True): the paper is student-facing, so NO answers leak —
    the correct/incorrect column is blank and the Solution row is empty regardless
    of what the question JSON carries. The correct answer + worked solution live
    ONLY on the separate answer-key / solution files. Set False only for an
    explicit teacher/master copy.
    """
    opts = q.get("options") or []
    n_opts = max(2, len(opts))          # at least the 2 label rows the layout needs
    # 4 fixed rows historically (Question, Type, then 4 Option rows, Solution,
    # Marks). Size the Option block to the ACTUAL option count so a 5th/6th option
    # is never silently dropped and an 'e'/'f' answer still lands on a real row.
    tbl = doc.add_table(rows=4 + n_opts, cols=3)
    _apply_tblPr(tbl)
    _apply_tblGrid(tbl)

    rows = tbl.rows
    sol_row = 2 + n_opts                # Solution row index (after the option rows)
    marks_row = sol_row + 1

    # Row 0: Question | <stem runs, colspan 2>
    _fill_label(rows[0].cells[0], "Question")
    stem_width = COL1_W + COL2_W
    _apply_tcPr(rows[0].cells[1]._tc, stem_width, span=2)
    tc2 = rows[0].cells[1]._tc.getnext()
    if tc2 is not None:
        tc2.getparent().remove(tc2)
    para = rows[0].cells[1].paragraphs[0]
    _fmt_para(para)
    for text, font in _runs(q["stem"]):
        _add_run(para, text, bold=True, font=font)

    # Row 1: Type | multiple_choice (colspan 2)
    _fill_label(rows[1].cells[0], "Type")
    _fill_plain(rows[1].cells[1], "multiple_choice",
                width=COL1_W + COL2_W, span=2, bold=False)

    # Which option (0-based) is the correct one, if the pipeline marked it from
    # the institute DB. Blank/absent answer -> no option marked correct (all
    # "incorrect" — the teacher fills it in).
    # NOTE: must test LIST membership, not `ans in "abcdef"`. A blank answer ""
    # is a SUBSTRING of "abcdef" (`"" in "abcdef"` is True), so the old check
    # let a blank through and `"abcdef".index("")` returned 0 -> Option A was
    # falsely marked "correct" on EVERY unanswered question (a fabricated key).
    # Which option (0-based) is the correct one. correct_idx is computed over all
    # SIX labels a..f so an 'e'/'f' key resolves — and because the option block is
    # now sized to len(options), that index always maps to a real rendered row.
    # NOTE: must test LIST membership, not `ans in "abcdef"`. A blank answer ""
    # is a SUBSTRING of "abcdef", so the old check let a blank through and
    # `.index("")` returned 0 -> Option A falsely "correct" on every blank.
    ans = str(q.get("answer", "")).strip().lower()
    correct_idx = list("abcdef").index(ans) if ans in list("abcdef") else -1

    # Option rows. On the STUDENT-SAFE paper the correct/incorrect column is ALWAYS
    # blank (no answer leaks); the answer + solution live only on the separate
    # answer-key / solution files. Only a teacher/master copy (student_safe=False)
    # prints the marking.
    has_answer = correct_idx >= 0
    for i, opt in enumerate(opts):
        row = rows[2 + i]
        _fill_label(row.cells[0], "Option")
        _fill_runs(row.cells[1], _runs(opt), width=COL1_W, bold=False)
        if student_safe:
            mark = ""
        elif has_answer:
            mark = "correct" if i == correct_idx else "incorrect"
        else:
            mark = ""  # no key -> blank, teacher fills it in
        _fill_plain(row.cells[2], mark, width=COL2_W, bold=False)

    # Solution row — colspan 2. Blank on the student-safe paper (the worked solution
    # is a separate teacher file); only a teacher copy embeds it.
    _fill_label(rows[sol_row].cells[0], "Solution")
    solution = "" if student_safe else str(q.get("solution", "")).strip()
    if solution:
        _apply_tcPr(rows[sol_row].cells[1]._tc, COL1_W + COL2_W, span=2)
        tc2 = rows[sol_row].cells[1]._tc.getnext()
        if tc2 is not None:
            tc2.getparent().remove(tc2)
        para = rows[sol_row].cells[1].paragraphs[0]
        _fmt_para(para)
        for text, font in _runs(solution):
            _add_run(para, text, bold=False, font=font)
    else:
        _fill_plain(rows[sol_row].cells[1], "",
                    width=COL1_W + COL2_W, span=2, bold=False)

    # Marks row
    _fill_label(rows[marks_row].cells[0], "Marks")
    _fill_plain(rows[marks_row].cells[1], "1",    width=COL1_W, bold=False)
    _fill_plain(rows[marks_row].cells[2], "0.25", width=COL2_W, bold=False)

    # Keep the whole question table on one page — MUST run AFTER all cells are
    # filled, because the fill/colspan helpers replace paragraphs (which would
    # drop a keepNext applied earlier).
    _no_split_rows(tbl)

    # Blank separator paragraph
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after  = Pt(0)
    sep.paragraph_format.space_before = Pt(0)


def build(src: Path, out_path: Path, font: str = DEFAULT_FONT,
          student_safe: bool = True):
    """Pipeline builder: universal questions JSON -> format-2 table .docx.

    student_safe (default True): produce a STUDENT paper with no answers/solutions
    embedded (the format-2 table would otherwise carry the correct-answer marking
    and worked solution in its own cells, and the file is named like a plain paper
    so a student could download it). The key + solutions ship as separate files."""
    global _FONT
    _FONT = font  # _runs() reads this via render_runs()
    data = json.loads(Path(src).read_text(encoding="utf-8"))
    questions = data["questions"]
    doc = Document()

    # Remove default empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    for q in questions:
        # skip a malformed question rather than crash the whole paper build
        if not str(q.get("stem", "")).strip() or not (q.get("options") or []):
            print(f"  [format2] skipping question with no stem/options: "
                  f"n={q.get('n')}", file=sys.stderr)
            continue
        add_question_table(doc, q, student_safe=student_safe)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"OK: {out_path}  ({len(questions)} questions)")


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    if len(sys.argv) < 2:
        print("usage: python build_paper_format2.py <questions.json> [out.docx]")
        sys.exit(1)
    src = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else (
        OUT_DIR / "Practice Paper - Target Academy (Format 2).docx"
    )
    build(src, out)
