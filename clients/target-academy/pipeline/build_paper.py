"""Build a question paper .docx in the institute's REAL template.

Approach: open one of the client's actual papers as the base document — this
inherits page geometry, the rectangular page borders (w:pgBorders), and the
1-col/2-col section structure — strip its content, then compose:

  page 1   : their cover image (from PPT FRONT.pptx), full page, section 0 (1 col)
  pages 2+ : questions in two columns (section 1, bordered), Kruti Dev 010 @ 14pt
  last page: उत्तरमाला (new single-column section)
  all pages: institute seal as a faded behind-text watermark (header anchor)

Text path: Unicode question JSON -> unicode_to_krutidev_runs() -> Kruti runs
(Kruti Dev 010) + Latin runs (Times New Roman).

Usage: python build_paper.py [questions.json] [out.docx]
"""

import copy
import json
import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Pt
from PIL import Image

sys.path.insert(0, str(Path(__file__).parent))
from krutidev import DEFAULT_FONT, render_runs  # noqa: E402

BASE = Path(__file__).resolve().parents[1]
FRONT_PAGE_DOCX = BASE / "resources" / "front-page.docx"  # made by extract_front_page.py
LOGO_IMG = BASE / "templates" / "docx_image1.jpeg"
# The watermark is regenerated from LOGO_IMG on every build, so it must be
# written to a WRITABLE path. templates/ is mounted read-only in production
# (client IP, :ro) — writing watermark.png there raises OSError [Errno 30].
# Generate it into a temp file instead.

LATIN = "Times New Roman"

# The font choice for the current build(), set by build(). add_mixed() reads it
# via render_runs() so the Devanagari font + encoding follow the frontend's pick.
_FONT = DEFAULT_FONT
BODY_PT = 14

# usable area inside their margins (A4, 1.27cm margins)
PAGE_W_CM, PAGE_H_CM, MARGIN_CM = 21.0, 29.7, 1.27
# Questions sit in 2 columns. Figures must stay COMPACT so a question + its
# figure + options fit together without spilling across a page break — matching
# how diagrams look small in the institute's own papers.
STEM_FIG_W_CM = 4.8   # a question's own diagram (Q3-type)
STEM_FIG_H_CM = 4.0   # height cap so wide-short crops don't blow up
OPT_FIG_W_CM = 3.6    # each option diagram (Q60-type) — smaller, sits beside label
OPT_FIG_H_CM = 3.0

_IMG_BASE = BASE  # set per build() so JSON image paths resolve beside the JSON


def _img_path(rel):
    """Resolve a question's image path, CONFINED to the job's image base.

    image/option_images come from the API request body (free strings), so an
    absolute path ("/etc/passwd") or a traversal ("../../../secret") would let a
    caller embed an arbitrary server-side file into their output doc (exfiltration).
    Reject absolute paths and any ref that resolves outside _IMG_BASE."""
    base = Path(_IMG_BASE).resolve()
    p = Path(rel)
    if p.is_absolute():
        raise ValueError(f"image path must be relative to the job dir: {rel!r}")
    dest = (base / p).resolve()
    if base != dest and base not in dest.parents:
        raise ValueError(f"image path escapes the job dir: {rel!r}")
    return dest


def add_figure(doc, rel, max_w_cm=STEM_FIG_W_CM, max_h_cm=STEM_FIG_H_CM, para=None):
    """Embed a cropped diagram, scaled to fit WITHIN (max_w_cm x max_h_cm).

    Caps by both width and height so neither a tall nor a wide crop overruns the
    column. If `para` is given, the picture is added to it (inline beside a
    label); otherwise a new centered paragraph is created.
    """
    path = _img_path(rel)
    with Image.open(path) as im:
        pw, ph = im.size
    w_cm = max_w_cm
    h_cm = w_cm * ph / pw
    if h_cm > max_h_cm:                     # too tall → constrain by height
        h_cm, w_cm = max_h_cm, max_h_cm * pw / ph
    if para is None:
        para = tight(doc.add_paragraph(), 4)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Cm(w_cm))
    return para


def add_mixed(para, text, size=BODY_PT, bold=False):
    runs, deva_font = render_runs(text, _FONT)
    for run_text, is_latin in runs:
        r = para.add_run(run_text)
        r.font.name = LATIN if is_latin else deva_font
        r.font.size = Pt(size)
        r.font.bold = bold
    return para


def add_latin(para, text, size=BODY_PT, bold=False):
    r = para.add_run(text)
    r.font.name = LATIN
    r.font.size = Pt(size)
    r.font.bold = bold
    return para


def tight(para, after=2):
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.space_before = Pt(0)
    return para


def _keep_with_next(para):
    """Set w:keepNext so this paragraph stays on the same page/column as the
    paragraph that follows it."""
    para.paragraph_format.keep_with_next = True


def _keep_lines(para):
    """Set w:keepLines so the paragraph's own lines don't split across a break."""
    para.paragraph_format.keep_together = True


def keep_question_together(paras):
    """Glue all paragraphs of one question (stem + statements + options) so a
    page/column break never lands BETWEEN a question and its options — the
    'question here, options on the next page' bug. Every paragraph keeps its own
    lines; every paragraph except the last keeps with the next."""
    for i, p in enumerate(paras):
        if p is None:
            continue
        _keep_lines(p)
        if i < len(paras) - 1:
            _keep_with_next(p)


def make_watermark(logo_path: Path, out_path: Path, alpha=0.10):
    """Faint logo with TRUE alpha transparency (no background) — text stays
    readable regardless of how a renderer stacks the watermark vs text."""
    with Image.open(logo_path) as src:  # context manager: no leaked file handle
        logo = src.convert("RGBA")
    a = logo.split()[3].point(lambda v: int(v * alpha))
    logo.putalpha(a)
    logo.save(out_path, "PNG")


def add_watermark_header(section, image_path: Path, width_cm=11.0):
    """Centered behind-text seal on every page of this section onward."""
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0]
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    inline = run._element.find(qn("w:drawing")).find(qn("wp:inline"))
    drawing = inline.getparent()
    anchor = drawing.makeelement(qn("wp:anchor"), {
        "distT": "0", "distB": "0", "distL": "114300", "distR": "114300",
        "simplePos": "0", "relativeHeight": "251658240", "behindDoc": "1",
        "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
    })
    simple_pos = anchor.makeelement(qn("wp:simplePos"), {"x": "0", "y": "0"})
    anchor.append(simple_pos)
    for axis, tag in (("positionH", "wp:positionH"), ("positionV", "wp:positionV")):
        pos = anchor.makeelement(qn(tag), {"relativeFrom": "page"})
        align = pos.makeelement(qn("wp:align"), {})
        align.text = "center"
        pos.append(align)
        anchor.append(pos)
    for child_tag in ("wp:extent", "wp:docPr", "a:graphic"):
        el = inline.find(qn(child_tag))
        anchor.append(copy.deepcopy(el))
        if child_tag == "wp:extent":
            anchor.append(anchor.makeelement(qn("wp:wrapNone"), {}))
    drawing.replace(inline, anchor)


def set_columns(section, num, space_twips=567):
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def _lock_table_widths(tbl, widths):
    """Force fixed column widths that Word/LibreOffice cannot autofit away.
    python-docx's cell.width alone is advisory — inside a narrow page column
    the layout engine renegotiates and can collapse a cell to zero. We set
    tblLayout=fixed and stamp an explicit <w:tcW> (in twips) on every cell,
    plus a <w:tblGrid> matching the widths, so the grid is authoritative."""
    tbl.autofit = False
    tbl.allow_autofit = False
    total_twips = 0
    for w in widths:
        total_twips += int(w.twips)
    tblPr = tbl._tbl.tblPr
    # fixed layout
    layout = tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"})
    tblPr.append(layout)
    # authoritative grid
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc in list(grid):
            grid.remove(gc)
        for w in widths:
            col = grid.makeelement(qn("w:gridCol"), {qn("w:w"): str(int(w.twips))})
            grid.append(col)
    # per-cell width
    for row in tbl.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = tcPr.makeelement(qn("w:tcW"),
                                   {qn("w:w"): str(int(w.twips)), qn("w:type"): "dxa"})
            tcPr.append(tcW)


def _opt_labels(q) -> list[str]:
    """One label per option, generated so we never index past the end. Questions
    can legitimately have 2-6+ options; a fixed list raised IndexError on more."""
    n = max(len(q.get("options") or []), len(q.get("option_images") or []), 5)
    return [f"({chr(ord('a') + i)})" if i < 26 else f"({i + 1})" for i in range(n)]


def add_questions(doc, questions):
    first_q_para = None
    for q in questions:
        labels = _opt_labels(q)
        # collect every paragraph this question produces so we can glue them
        # together (no page/column break between a stem and its options).
        q_paras = []
        p = tight(doc.add_paragraph())
        q_paras.append(p)
        if first_q_para is None:
            first_q_para = p
        add_mixed(p, f"{q['n']}. {q['stem']}")
        for i, st in enumerate(q.get("statements") or [], start=1):
            sp = tight(doc.add_paragraph())
            add_mixed(sp, f"{i}. {st}")
            q_paras.append(sp)
        match_rows = q.get("match") or []
        if match_rows:
            # सूची-I / सूची-II as a REAL two-column table. Inline "left - right"
            # paragraphs wrapped into soup (client feedback 2026-07-14). This
            # table lives INSIDE the paper's narrow 2-column section (~8.5cm
            # per page-column), so the two cells MUST sum to less than that or
            # Word collapses the second cell off-view (सूची-II vanished on the
            # first prod paper). Cells 4.7 + 3.4 = 8.1cm, locked at the XML
            # level (fixed layout + explicit tcW) so autofit can't override.
            tbl = doc.add_table(rows=1 + len(match_rows), cols=2)
            try:
                tbl.style = doc.styles["Normal Table"]  # borderless everywhere
            except KeyError:
                pass
            _lock_table_widths(tbl, (Cm(4.7), Cm(3.4)))
            for cell, title in zip(tbl.rows[0].cells, ("सूची-I", "सूची-II")):
                hp = tight(cell.paragraphs[0])
                add_mixed(hp, title, bold=True)
                q_paras.append(hp)
            for r, (left, right) in enumerate(match_rows, start=1):
                for cell, text in zip(tbl.rows[r].cells, (left, right)):
                    cp = tight(cell.paragraphs[0])
                    add_mixed(cp, str(text))
                    q_paras.append(cp)
        if q.get("lead_in"):
            lp = tight(doc.add_paragraph())
            add_mixed(lp, q["lead_in"])
            q_paras.append(lp)
        # the question's own diagram (Q3-type), embedded verbatim under the stem
        if q.get("image"):
            add_figure(doc, q["image"])
        # option diagrams (Q60-type): each (a)-(d) is a cropped figure, kept
        # compact and inline beside its label so all four fit in the column
        if q.get("option_images"):
            for i, rel in enumerate(q["option_images"]):
                op = tight(doc.add_paragraph(), 2)
                add_latin(op, labels[i] + " ")
                add_figure(doc, rel, max_w_cm=OPT_FIG_W_CM, max_h_cm=OPT_FIG_H_CM, para=op)
                q_paras.append(op)
        else:
            opts = q["options"]
            if not q.get("long_options") and len(opts) == 4 and all(len(o) <= 20 for o in opts):
                for i in (0, 2):
                    op = tight(doc.add_paragraph())
                    add_latin(op, labels[i] + " ")
                    add_mixed(op, opts[i])
                    add_latin(op, "\t\t")
                    add_latin(op, labels[i + 1] + " ")
                    add_mixed(op, opts[i + 1])
                    q_paras.append(op)
            else:
                for i, o in enumerate(opts):
                    op = tight(doc.add_paragraph())
                    add_latin(op, labels[i] + " ")
                    add_mixed(op, o)
                    q_paras.append(op)
        # keep the whole question (stem -> options) on one page/column
        keep_question_together(q_paras)
        tight(doc.add_paragraph(), 6)
    return first_q_para


def add_answer_key(doc, questions):
    key = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_columns(key, 1)
    p = tight(doc.add_paragraph(), 6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_mixed(p, "उत्तरमाला", size=16, bold=True)
    row = []
    for q in questions:
        # tolerate a missing/blank answer (unmarked question / no key) — render a
        # clean dash instead of "5. ()" or KeyError-ing. Matches build_answer_key,
        # build_deck, build_solution which all use .get("answer", ""). .lower() so
        # an uppercase "A" prints "(a)" — consistent with the standalone answer
        # key (build_answer_key lowercases too), not "(A)" vs "(a)".
        ans = str(q.get("answer", "")).strip().lower()
        row.append(f"{q.get('n', '')}. ({ans or '—'})")
        if len(row) == 5:
            add_latin(tight(doc.add_paragraph(), 4), "      ".join(row), size=12)
            row = []
    if row:
        add_latin(tight(doc.add_paragraph(), 4), "      ".join(row), size=12)


def build(questions_path: Path, out_path: Path, font: str = DEFAULT_FONT):
    global _IMG_BASE, _FONT
    _FONT = font  # add_mixed() reads this via render_runs()
    _IMG_BASE = questions_path.resolve().parent  # crops resolve beside the JSON
    data = json.loads(questions_path.read_text(encoding="utf-8"))
    doc = Document(str(FRONT_PAGE_DOCX))  # their real front page, original sections
    # regenerate the watermark into a writable temp file (templates/ is :ro)
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as _wm:
        watermark_img = Path(_wm.name)
    try:
        make_watermark(LOGO_IMG, watermark_img)  # always regenerate (recipe may change)
        add_watermark_header(doc.sections[-1], watermark_img)
        first_q_para = add_questions(doc, data["questions"])
        # questions must start on a fresh page after the front matter
        ppr = first_q_para._p.get_or_add_pPr()
        ppr.insert(0, ppr.makeelement(qn("w:pageBreakBefore"), {}))
        add_answer_key(doc, data["questions"])
        # Referral QR on the last page (opt-in via REFERRAL_QR=1). Every paper a
        # student receives becomes a self-registration point for the referral
        # program, with zero owner effort. Best-effort: a QR failure must never
        # break the paper build.
        if os.environ.get("REFERRAL_QR", "").strip() == "1":
            try:
                from referral_qr import add_referral_block_docx
                add_referral_block_docx(doc)
            except Exception as e:  # noqa: BLE001
                print(f"WARN: referral QR skipped: {e}")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
    finally:
        # The watermark PNG is already embedded in the docx; remove the temp
        # file so it does not leak one PNG per build.
        watermark_img.unlink(missing_ok=True)
    print(f"OK: {out_path}")


def verify(out_path: Path):
    from ingest import paragraph_text

    doc = Document(str(out_path))
    texts = [paragraph_text(p).strip() for p in doc.paragraphs if paragraph_text(p).strip()]
    dev = sum(1 for t in texts for c in t if 0x0900 <= ord(c) <= 0x097F)
    borders = doc.element.body.xml.count("pgBorders")
    print(f"verify: {len(texts)} paragraphs, {dev} Devanagari chars round-tripped, "
          f"pgBorders={borders}, sections={len(doc.sections)}")
    print("sample:", texts[0][:80])


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    q = Path(sys.argv[1]) if len(sys.argv) > 1 else BASE / "review" / "spike3-questions.json"
    if len(sys.argv) > 2:
        out = Path(sys.argv[2])
    else:
        # client-facing filename from the questions JSON ("filename" key)
        data = json.loads(q.read_text(encoding="utf-8"))
        out = BASE / "review" / data.get("filename", "Practice Paper - Target Academy.docx")
    build(q, out)
    verify(out)
