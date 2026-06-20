"""One-off: extract template anatomy from the client's real papers.

Reports per file: page/section setup (size, margins, columns), header/footer
text, the first paragraphs (title block) with fonts/sizes, option-line layout,
and the tail (answer-key format). Feeds brain/format-conventions.md + paper.py.
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Twips

sys.path.insert(0, str(Path(__file__).parent))
from ingest import paragraph_text  # noqa: E402

RES = Path(__file__).resolve().parents[1] / "resources"


def emu_to_cm(v):
    return round(v / 360000, 2) if v else None


def describe(path: Path):
    print(f"\n{'='*72}\n{path.name}\n{'='*72}")
    doc = Document(str(path))

    for si, sec in enumerate(doc.sections):
        cols = sec._sectPr.xpath("./w:cols")
        ncols = cols[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num") if cols else None
        space = cols[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space") if cols else None
        print(f"section {si}: page {emu_to_cm(sec.page_width)}x{emu_to_cm(sec.page_height)}cm "
              f"margins L{emu_to_cm(sec.left_margin)} R{emu_to_cm(sec.right_margin)} "
              f"T{emu_to_cm(sec.top_margin)} B{emu_to_cm(sec.bottom_margin)} "
              f"| cols={ncols or 1} space={space and Twips(int(space)).cm:.2}cm" if space else
              f"section {si}: page {emu_to_cm(sec.page_width)}x{emu_to_cm(sec.page_height)}cm "
              f"margins L{emu_to_cm(sec.left_margin)} R{emu_to_cm(sec.right_margin)} | cols={ncols or 1}")
        hdr = " / ".join(p.text.strip() for p in sec.header.paragraphs if p.text.strip())
        ftr = " / ".join(p.text.strip() for p in sec.footer.paragraphs if p.text.strip())
        if hdr:
            print(f"  header: {hdr[:120]!r}")
        if ftr:
            print(f"  footer: {ftr[:120]!r}")

    print("\n--- first 12 paragraphs (title block) ---")
    for p in doc.paragraphs[:12]:
        t = paragraph_text(p).strip()
        if not t:
            continue
        fonts = {(r.font.name, r.font.size.pt if r.font.size else None, r.font.bold) for r in p.runs if r.text.strip()}
        print(f"  [{p.alignment} | {p.style.name} | {fonts}] {t[:90]!r}")

    print("\n--- last 10 non-empty paragraphs (answer key?) ---")
    tail = [p for p in doc.paragraphs if paragraph_text(p).strip()][-10:]
    for p in tail:
        print(f"  {paragraph_text(p).strip()[:100]!r}")

    print(f"\ntables: {len(doc.tables)}")
    if doc.tables:
        t0 = doc.tables[-1]
        print(f"  last table: {len(t0.rows)} rows x {len(t0.columns)} cols")
        for row in t0.rows[:3]:
            print("   | " + " | ".join(paragraph_text(p).strip() for c in row.cells for p in c.paragraphs)[:110])


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    for name in ["TARGET SERIES 01.docx", "TARGET SERIES 02.docx"]:
        describe(RES / name)
