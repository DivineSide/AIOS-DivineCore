"""Build the TEACHER solution doc (.docx) from a solved questions JSON.

Private to the teacher (NOT shown in class, NOT branded publicly). Follows the
same visual standard as the practice paper: opens front-page.docx to inherit the
rectangular page border (w:pgBorders), applies the faded watermark logo, and uses
Kruti Dev 010 / Times New Roman via unicode_to_krutidev_runs — exactly as
build_paper.py does.

Per question layout (matches the WhatsApp solution format):
    Q1. <stem>                          — Kruti Dev, bold
    (A) <opt>  (B) <opt> ...            — Kruti Dev
    उत्तर: (C) <correct option text>     — bold, green
    विस्तृत व्याख्या: <explanation>      — Kruti Dev, label bold

Source provenance and ⚠ flags live HERE ONLY — never on the branded paper/deck.

Usage: python build_solution.py [questions.json] [out.docx]
"""

import copy
import json
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

sys.path.insert(0, str(Path(__file__).parent))
from krutidev import DEFAULT_FONT, render_runs  # noqa: E402

# Font for the current build(), set by build(); add_mixed reads it via render_runs.
_FONT = DEFAULT_FONT

BASE         = Path(__file__).resolve().parents[1]
FRONT_PAGE_DOCX = BASE / "resources" / "front-page.docx"
LOGO_IMG     = BASE / "templates" / "docx_image1.jpeg"
# watermark is regenerated each build -> writable temp file, not the :ro
# templates mount (writing there raises OSError [Errno 30]).

LATIN  = "Times New Roman"
BODY_PT = 12
FIG_W_CM = 6.0      # single-column doc — can be a bit wider than 2-col paper

GREEN = RGBColor(0x1B, 0x6B, 0x2D)
GREY  = RGBColor(0x55, 0x55, 0x55)
RED   = RGBColor(0xC0, 0x2A, 0x2A)

LABELS     = ["a", "b", "c", "d", "e", "f"]
LABELS_UP  = ["A", "B", "C", "D", "E", "F"]

_IMG_BASE = BASE


def _img_path(rel):
    p = Path(rel)
    return p if p.is_absolute() else (_IMG_BASE / p)


# ── shared helpers (same pattern as build_paper.py) ──────────────────────────

def make_watermark(logo_path: Path, out_path: Path, alpha=0.10):
    logo = Image.open(logo_path).convert("RGBA")
    a = logo.split()[3].point(lambda v: int(v * alpha))
    logo.putalpha(a)
    logo.save(out_path, "PNG")


def add_watermark_header(section, image_path: Path, width_cm=11.0):
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0]
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    inline = run._element.find(qn("w:drawing")).find(qn("wp:inline"))
    drawing = inline.getparent()
    anchor = drawing.makeelement(qn("wp:anchor"), {
        "distT": "0", "distB": "0", "distL": "114300", "distR": "114300",
        "simplePos": "0", "relativeHeight": "251658240", "behindDoc": "1",
        "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
    })
    simple_pos = anchor.makeelement(qn("wp:simplePos"), {"x": "0", "y": "0"})
    anchor.append(simple_pos)
    for axis, tag in (("positionH", "wp:positionH"), ("positionV", "wp:positionV")):
        pos = anchor.makeelement(qn(tag), {"relativeFrom": "page"})
        align = pos.makeelement(qn("wp:align"), {})
        align.text = "center"
        pos.append(align)
        anchor.append(pos)
    for child_tag in ("wp:extent", "wp:docPr", "a:graphic"):
        el = inline.find(qn(child_tag))
        anchor.append(copy.deepcopy(el))
        if child_tag == "wp:extent":
            anchor.append(anchor.makeelement(qn("wp:wrapNone"), {}))
    drawing.replace(inline, anchor)


def tight(p, after=2, before=0):
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    return p


def add_mixed(para, text, size=BODY_PT, bold=False, color=None):
    """Unicode string -> Devanagari + Times New Roman runs, font per _FONT."""
    runs, deva_font = render_runs(text, _FONT)
    for run_text, is_latin in runs:
        r = para.add_run(run_text)
        r.font.name  = LATIN if is_latin else deva_font
        r.font.size  = Pt(size)
        r.font.bold  = bold
        if color:
            r.font.color.rgb = color
    return para


def add_latin(para, text, size=BODY_PT, bold=False, color=None):
    r = para.add_run(text)
    r.font.name  = LATIN
    r.font.size  = Pt(size)
    r.font.bold  = bold
    if color:
        r.font.color.rgb = color
    return para


# ── per-question renderer ─────────────────────────────────────────────────────

def _label_up(i: int) -> str:
    """Uppercase option label for any index (A, B, ... Z, then numbers)."""
    return chr(ord("A") + i) if i < 26 else str(i + 1)


def add_question(doc, q, official: bool):
    # Tolerate a missing/odd answer (imported papers often have no printed key).
    # -1 means "no correct option to highlight" rather than a hard crash.
    ans = str(q.get("answer", "")).strip().lower()
    ans_idx = LABELS.index(ans) if ans in LABELS else -1

    # Q1. stem
    qp = tight(doc.add_paragraph(), after=2, before=8)
    add_mixed(qp, f"{q['n']}. {q['stem']}", bold=True)

    # sub-statements (assertion style)
    for i, st in enumerate(q.get("statements") or [], start=1):
        add_mixed(tight(doc.add_paragraph(), 1), f"   {i}. {st}")

    # match pairs
    for left_t, right_t in q.get("match") or []:
        mp = tight(doc.add_paragraph(), 1)
        add_mixed(mp, f"   {left_t}")
        add_latin(mp, "  —  ")
        add_mixed(mp, right_t)

    if q.get("lead_in"):
        add_mixed(tight(doc.add_paragraph(), 2), f"   {q['lead_in']}")

    # question diagram
    if q.get("image"):
        fp = tight(doc.add_paragraph(), 4)
        fp.add_run().add_picture(str(_img_path(q["image"])), width=Cm(FIG_W_CM))

    # options
    if q.get("option_images"):
        for i, rel in enumerate(q["option_images"]):
            correct = i == ans_idx
            op = tight(doc.add_paragraph(), 1)
            add_latin(op, f"   ({_label_up(i)}) ",
                      bold=correct, color=GREEN if correct else None)
            if correct:
                add_latin(op, "✓ ", bold=True, color=GREEN)
            op.add_run().add_picture(str(_img_path(rel)), width=Cm(FIG_W_CM - 2.0))
    else:
        # q.get("options") or [] — a question with a missing/null options key
        # (manual-import or a partial extract) must not crash the whole build;
        # solutions_task/answer_key_task call this directly, bypassing the
        # validate() drop-gate that the build/full pipeline runs.
        for i, opt in enumerate(q.get("options") or []):
            correct = i == ans_idx
            op = tight(doc.add_paragraph(), 1)
            mark = " ✓" if correct else ""
            add_mixed(op, f"   ({_label_up(i)}) {opt}{mark}",
                      bold=correct, color=GREEN if correct else None)

    # उत्तर line
    ansp = tight(doc.add_paragraph(), after=1, before=2)
    add_mixed(ansp, "उत्तर: ", bold=True, color=GREEN)
    opts = q.get("options") or []
    if ans_idx < 0:
        # no printed answer key in the source paper — say so rather than guess
        add_mixed(ansp, "(not provided)", bold=True, color=GREEN)
    else:
        ans_txt = f"({_label_up(ans_idx)})"
        if ans_idx < len(opts):
            ans_txt += f" {opts[ans_idx]}"
        add_mixed(ansp, ans_txt, bold=True, color=GREEN)

    # विस्तृत व्याख्या (solution) or संकेत (short reason)
    explanation = q.get("solution") or q.get("reason")
    if explanation:
        label = "विस्तृत व्याख्या: " if q.get("solution") else "संकेत: "
        ep = tight(doc.add_paragraph(), after=2)
        add_mixed(ep, label, bold=True)
        add_mixed(ep, explanation)

    # Source provenance + flag — TEACHER DOC ONLY
    srcs = [s.lower() for s in (q.get("sources") or [])]
    if official:
        sp = tight(doc.add_paragraph(), 2)
        add_mixed(sp, "✔ स्रोत: ", size=9, bold=True, color=GREEN)
        add_mixed(sp, "आधिकारिक उत्तर कुंजी (UKSSSC)", size=9, color=GREY)
    elif srcs:
        sp = tight(doc.add_paragraph(), 2)
        add_mixed(sp, "✔ स्रोत: ", size=9, bold=True, color=GREEN)
        # add_mixed (not add_latin): a source book name can contain Devanagari
        # (book_chunks.book_name is a filename, e.g. "KIRAN सामान्य हिन्दी.pdf"),
        # and add_latin forces Times New Roman which has no Devanagari glyphs ->
        # tofu boxes. add_mixed gives non-Latin runs a real Devanagari font.
        add_mixed(sp, " / ".join(srcs), size=9, color=GREY)
    elif q.get("flag"):
        sp = tight(doc.add_paragraph(), 2)
        add_mixed(sp, "⚠ जाँच आवश्यक — ", size=10, bold=True, color=RED)
        add_mixed(sp, q["flag"], size=10, color=RED)


# ── build entry point ─────────────────────────────────────────────────────────

def build(questions_path: Path, out_path: Path, font: str = DEFAULT_FONT):
    global _IMG_BASE, _FONT
    _FONT = font  # add_mixed() reads this via render_runs()
    _IMG_BASE = questions_path.resolve().parent

    data     = json.loads(questions_path.read_text(encoding="utf-8"))
    official = data.get("answer_source") == "official_key"

    # Open the branded template — inherits pgBorders, page geometry, section def
    doc = Document(str(FRONT_PAGE_DOCX))

    # Strip every paragraph and table that came with the template content
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)

    # Watermark on first (only) section
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as _wm:
        watermark_img = Path(_wm.name)
    make_watermark(LOGO_IMG, watermark_img)
    add_watermark_header(doc.sections[0], watermark_img)

    # Title
    title = tight(doc.add_paragraph(), after=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_latin(title,
              data.get("solution_title", "समाधान — शिक्षक हेतु"),
              size=16, bold=True)

    sub = tight(doc.add_paragraph(), after=10)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_latin(sub,
              data.get("solution_subtitle", "केवल शिक्षक के लिए — कक्षा में प्रदर्शित न करें"),
              size=10, color=GREY)

    for q in data["questions"]:
        add_question(doc, q, official)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    n        = len(data["questions"])
    with_sol = sum(1 for q in data["questions"] if q.get("solution"))
    print(f"OK: {out_path} ({n} questions, {with_sol} with worked solutions)")


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    q = Path(sys.argv[1]) if len(sys.argv) > 1 else BASE / "review" / "spike3-questions.json"
    data = json.loads(q.read_text(encoding="utf-8"))
    default = data.get("solution_filename") or (
        Path(data.get("filename", "Solution.docx")).stem + " - Solution (Teacher).docx")
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else BASE / "review" / default
    build(q, out)
