"""Build a branded ANSWER-KEY PDF (उत्तरमाला) from a questions JSON.

Generates a properly branded answer key in the institute's real template:
  - Same front-page .docx base (page borders, margins, section geometry)
  - Watermark logo behind text
  - Kruti Dev 010 font for Devanagari (matches the practice paper exactly)
  - Answer grid: 5 columns, "N. (x)" format

Renders as .docx first (using build_paper.py helpers) then converts to PDF
via win32com (Word on Windows) so the output is print/WhatsApp-ready.

Usage: python build_answer_key.py [questions.json] [out.pdf]
"""

import json
import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).parent))
from krutidev import DEFAULT_FONT, render_runs  # noqa: E402

BASE = Path(__file__).resolve().parents[1]
FRONT_PAGE_DOCX = BASE / "resources" / "front-page.docx"
LOGO_IMG = BASE / "templates" / "docx_image1.jpeg"
# watermark is regenerated each build -> write to a writable temp file, not the
# read-only (:ro) templates mount (which raises OSError [Errno 30]).

LATIN = "Times New Roman"
_FONT = DEFAULT_FONT  # set by build(); add_mixed() reads it via render_runs()


def add_mixed(para, text, size=14, bold=False):
    runs, deva_font = render_runs(text, _FONT)
    for run_text, is_latin in runs:
        r = para.add_run(run_text)
        r.font.name = LATIN if is_latin else deva_font
        r.font.size = Pt(size)
        r.font.bold = bold
    return para


def add_latin(para, text, size=14, bold=False):
    r = para.add_run(text)
    r.font.name = LATIN
    r.font.size = Pt(size)
    r.font.bold = bold
    return para


def tight(para, after=2):
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.space_before = Pt(0)
    return para


def _make_watermark(logo_path: Path, out_path: Path, alpha: float = 0.10):
    from PIL import Image
    logo = Image.open(logo_path).convert("RGBA")
    a = logo.split()[3].point(lambda v: int(v * alpha))
    logo.putalpha(a)
    logo.save(out_path, "PNG")


def _add_watermark_header(section, image_path: Path, width_cm: float = 11.0):
    import copy
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0]
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    inline = run._element.find(qn("w:drawing")).find(qn("wp:inline"))
    drawing = inline.getparent()
    anchor = drawing.makeelement(qn("wp:anchor"), {
        "distT": "0", "distB": "0", "distL": "114300", "distR": "114300",
        "simplePos": "0", "relativeHeight": "251658240", "behindDoc": "1",
        "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
    })
    simple_pos = anchor.makeelement(qn("wp:simplePos"), {"x": "0", "y": "0"})
    anchor.append(simple_pos)
    for tag in ("wp:positionH", "wp:positionV"):
        pos = anchor.makeelement(qn(tag), {"relativeFrom": "page"})
        align = pos.makeelement(qn("wp:align"), {})
        align.text = "center"
        pos.append(align)
        anchor.append(pos)
    for child_tag in ("wp:extent", "wp:docPr", "a:graphic"):
        el = inline.find(qn(child_tag))
        anchor.append(copy.deepcopy(el))
        if child_tag == "wp:extent":
            anchor.append(anchor.makeelement(qn("wp:wrapNone"), {}))
    drawing.replace(inline, anchor)


def _set_columns(section, num, space_twips=567):
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def _build_docx(questions_path: Path, docx_path: Path):
    data = json.loads(questions_path.read_text(encoding="utf-8"))
    questions = data["questions"]

    # Start from the real template (inherits page borders, margins, geometry)
    # but skip the front page — open it, clear all content, keep section/border XML
    doc = Document(str(FRONT_PAGE_DOCX))

    # Remove every paragraph and table from the body so we start blank
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)

    # Force single column on the (only) section so the grid fills full width
    _set_columns(doc.sections[-1], 1)

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as _wm:
        watermark_img = Path(_wm.name)
    _make_watermark(LOGO_IMG, watermark_img)
    _add_watermark_header(doc.sections[-1], watermark_img)

    # Title
    title_text = data.get("answer_key_title") or "उत्तरमाला"
    subtitle_text = (data.get("answer_key_subtitle")
                     or f"{data.get('title_latin', 'TARGET ACADEMY')}  —  उत्तर कुंजी")

    p = tight(doc.add_paragraph(), 6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_mixed(p, title_text, size=20, bold=True)

    p2 = tight(doc.add_paragraph(), 16)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_mixed(p2, subtitle_text, size=13)

    # Answer grid as a table — 5 columns, auto-width, centred
    # Each cell: "N. (x)", Times New Roman 13pt, centred
    COLS = 5
    rows_needed = -(-len(questions) // COLS)  # ceil division
    from docx.shared import Inches
    tbl = doc.add_table(rows=rows_needed, cols=COLS)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore

    # Remove all borders so it looks clean (borderless table)
    from docx.oxml import OxmlElement
    tbl_pr = tbl._tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    for idx, q in enumerate(questions):
        r, c = divmod(idx, COLS)
        cell = tbl.cell(r, c)
        cell_p = cell.paragraphs[0]
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_p.paragraph_format.space_before = Pt(6)
        cell_p.paragraph_format.space_after = Pt(6)
        add_latin(cell_p, f"{q['n']}. ({str(q['answer']).lower()})", size=13)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def _docx_to_pdf(docx_path: Path, pdf_path: Path):
    """Convert .docx -> .pdf. Word via win32com on Windows (dev), LibreOffice
    headless on Linux (server). Tries Word first, falls back to soffice."""
    if sys.platform == "win32":
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(str(docx_path.resolve()))
                doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)  # 17 = wdFormatPDF
                doc.Close()
            finally:
                word.Quit()
            return
        except Exception:
            pass  # fall through to LibreOffice if Word is unavailable
    _docx_to_pdf_libreoffice(docx_path, pdf_path)


def _docx_to_pdf_libreoffice(docx_path: Path, pdf_path: Path):
    """Server path: LibreOffice headless. Requires `libreoffice`/`soffice` on PATH
    plus the Kruti Dev / Nirmala UI fonts installed system-wide."""
    import shutil
    import subprocess

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) not found on PATH.")
    out_dir = pdf_path.resolve().parent
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir",
         str(out_dir), str(docx_path.resolve())],
        check=True, capture_output=True, timeout=120,
    )
    # LibreOffice names the output <docx-stem>.pdf — rename to the requested name.
    produced = out_dir / (docx_path.stem + ".pdf")
    if produced != pdf_path and produced.exists():
        produced.replace(pdf_path)


def build(questions_path: Path, out_path: Path, font: str = DEFAULT_FONT):
    global _FONT
    _FONT = font  # add_mixed() reads this via render_runs()
    data = json.loads(questions_path.read_text(encoding="utf-8"))
    n = len(data["questions"])

    # Build .docx first (same template pipeline)
    tmp_docx = out_path.with_suffix(".docx")
    _build_docx(questions_path, tmp_docx)

    # Convert to PDF
    try:
        _docx_to_pdf(tmp_docx, out_path)
        # Remove the intermediate .docx (it was just a stepping stone)
        tmp_docx.unlink(missing_ok=True)
    except Exception as e:
        # If Word isn't available, leave the .docx and rename out_path to .docx
        out_path = tmp_docx
        print(f"  (PDF conversion unavailable: {e} — saved as .docx instead)")

    print(f"OK: {out_path} ({n} answers, 5 cols)")


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    q = Path(sys.argv[1]) if len(sys.argv) > 1 else (
        BASE / "review" / "input" / "live-14jun-paper.json")
    data = json.loads(q.read_text(encoding="utf-8"))
    default = data.get("answer_key_filename") or (
        Path(data.get("filename", "Answer Key.docx")).stem + " - Answer Key.pdf")
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else (
        BASE / "review" / "output" / default)
    build(q, out)
